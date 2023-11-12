from django.contrib.auth.models import User
from django.http import HttpResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.urls import reverse_lazy
from docx.shared import Inches
from docx.shared import Pt
from .models import Task, SubTask, Activity, ActivityImages, Unit
from .forms import ActivityForm, PasswordChangingForm
from docx import Document
from django.db.models import Q
import datetime
from jalali_date import date2jalali, datetime2jalali
from django.contrib.auth.decorators import login_required
from django.contrib.auth.views import PasswordChangeView
from django.contrib.auth.forms import PasswordChangeForm


@login_required
def home(request):
    if request.user.is_staff:
        group = request.user.groups.all()[0].name
        task = Task.objects.all()
        sub_task = SubTask.objects.all()
        activities = Activity.objects.filter(name=group)
        full_name = request.user.get_full_name()

        for x in activities:
            x.shamsi = date2jalali(x.created)
            x.save()

        query = request.GET.get('query', '')
        if query:
            activities = activities.filter(
                Q(task__icontains=query) | Q(sub_task__icontains=query) | Q(name__icontains=query) | Q(
                    created__icontains=query) | Q(shamsi__icontains=query) | Q(str_user__icontains=query))
        context = {

            'task': task,
            'sub_task': sub_task,
            'activities': activities,
            'full_name': full_name,

        }

        return render(request, 'home.html', context)

    else:
        print("employee")
        task = Task.objects.all()
        sub_task = SubTask.objects.all()
        activities = Activity.objects.filter(user=request.user)
        full_name = request.user.get_full_name()

        for x in activities:
            x.shamsi = date2jalali(x.created)
            x.save()
        query = request.GET.get('query', '')
        if query:
            activities = activities.filter(
                Q(task__icontains=query) | Q(sub_task__icontains=query) | Q(name__icontains=query) | Q(
                    created__icontains=query) | Q(shamsi__icontains=query))
        context = {

            'task': task,
            'sub_task': sub_task,
            'activities': activities,
            'full_name': full_name,

        }

        return render(request, 'home.html', context)


@login_required
def create_activity(request):
    if request.method == 'POST':
        activity_form = ActivityForm(request.POST, request.FILES, request.user)
        if activity_form.is_valid:
            obj = activity_form.save(commit=False)
            obj.user = request.user
            obj.str_user = request.user.get_full_name()
            obj.save()
            p = activity_form.save()
            images = request.FILES.getlist("more_images")
            for i in images:
                ActivityImages.objects.create(activity=p, image=i)
            return redirect('home')
    else:
        activity_form = ActivityForm()

    tasks = Task.objects.all()
    sub_tasks = SubTask.objects.all()
    units = Unit.objects.all()
    full_name = request.user.get_full_name()
    my_group = request.user.groups.all()[0].name
    users = User.objects.filter(groups__name=my_group)
    for x in users:
        print(x.get_full_name())

    context = {
        'form': activity_form,
        'tasks': tasks,
        'sub_tasks': sub_tasks,
        'units': units,
        'full_name': full_name,
        'users': users,
    }

    return render(request, 'create_activity.html', context)


@login_required
def detail_activity(request, id):
    task = Task.objects.all()
    sub_task = SubTask.objects.all()
    activities = Activity.objects.all()
    activity = Activity.objects.get(pk=id)
    full_name = request.user.get_full_name()
    context = {

        'task': task,
        'sub_task': sub_task,
        'activities': activities,
        'activity': activity,
        'full_name': full_name,
    }

    return render(request, 'detail_activity.html', context)


@login_required
def delete_activity(request, id):
    activity = Activity.objects.filter(id=id)
    activity.delete()
    return redirect('home')


@login_required
def edit_activity(request, id):
    activity = Activity.objects.get(id=id)

    if request.method == 'POST':
        form = ActivityForm(request.POST, request.FILES, instance=activity)
        if form.is_valid():
            form.save()
            return redirect('home')

    else:
        form = ActivityForm(instance=activity)

    tasks = Task.objects.all()
    sub_tasks = SubTask.objects.all()
    units = Unit.objects.all()
    full_name = request.user.get_full_name()
    my_group = request.user.groups.all()[0].name
    users = User.objects.filter(groups__name=my_group)

    context = {
        'form': form,
        'tasks': tasks,
        'sub_tasks': sub_tasks,
        'units': units,
        'activity_image': activity.image,
        'full_name': full_name,
        'users': users,
    }

    return render(request, 'edit_activity.html', context)


@login_required
def export(request):
    if request.user.is_staff:
        group = request.user.groups.all()[0].name
        response = HttpResponse(content_type="application/ms-word")
        response["content-Disposition"] = 'attachment;filename=report' + str(datetime.datetime.now()) + '.docx'
        activities = Activity.objects.filter(name=group)
        activityimages = ActivityImages.objects.all()
        full_name = request.user.get_full_name()

        query = request.GET.get('query', '')
        if query:
            activities = activities.filter(
                Q(task__icontains=query) | Q(sub_task__icontains=query) | Q(name__icontains=query) | Q(
                    created__icontains=query) | Q(shamsi__icontains=query) | Q(str_user__icontains=query))

            document = Document()
            for x in activities:
                document.add_heading(x.name, 0)
                p = document.add_paragraph(str(x.str_user))
                p = document.add_paragraph(x.task)
                p.add_run("_")
                p.add_run(x.sub_task)
                p = document.add_paragraph("مدت زمان انجام فعالیت")
                p.add_run(str(x.duration))
                p.add_run("ساعت")
                if x.colleague:
                    p = document.add_paragraph("مشارکت همکار / همکاران:")
                    p = document.add_paragraph(x.colleague)
                p = document.add_paragraph(":تاریخ انجام فعالیت")
                p = document.add_paragraph(str(date2jalali(x.created)))

                document.add_picture(x.image, width=Inches(3.9))
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                for y in activityimages:
                    if y.activity == x:
                        run.add_picture(y.image, width=Inches(2.3))

                p = document.add_paragraph(":توضیحات")
                p = document.add_paragraph(x.description)
                document.add_page_break()
            document.save(response)
            return response

        return render(request, 'export.html', {'full_name': full_name})
    else:
        response = HttpResponse(content_type="application/ms-word")
        response["content-Disposition"] = 'attachment;filename=report' + str(datetime.datetime.now()) + '.docx'
        activities = Activity.objects.filter(str_user=str(request.user.get_full_name()))
        activityimages = ActivityImages.objects.all()
        full_name = request.user.get_full_name()

        query = request.GET.get('query', '')
        if query:
            activities = activities.filter(
                Q(task__icontains=query) | Q(sub_task__icontains=query) | Q(name__icontains=query) | Q(
                    created__icontains=query) | Q(shamsi__icontains=query) | Q(str_user__icontains=query))

            document = Document()
            for x in activities:
                document.add_heading(x.name, 0)
                p = document.add_paragraph(str(x.str_user))
                p = document.add_paragraph(x.task)
                p.add_run("_")
                p.add_run(x.sub_task)
                p = document.add_paragraph("مدت زمان انجام فعالیت")
                p.add_run(str(x.duration))
                p.add_run("ساعت")
                if x.colleague:
                    p = document.add_paragraph(":مشارکت همکار / همکاران")
                    p = document.add_paragraph(x.colleague)
                p = document.add_paragraph(":تاریخ انجام فعالیت")
                p = document.add_paragraph(str(date2jalali(x.created)))

                document.add_picture(x.image, width=Inches(3.9))
                paragraph = document.add_paragraph()
                run = paragraph.add_run()
                for y in activityimages:
                    if y.activity == x:
                        run.add_picture(y.image, width=Inches(2.3))

                p = document.add_paragraph(":توضیحات")
                p = document.add_paragraph(x.description)
                document.add_page_break()
            document.save(response)
            return response

        my_group = request.user.groups.all()[0].name
        users = User.objects.filter(groups__name=my_group)
        for x in users:
            print("x", x.get_full_name())
            print("my_group", my_group)
        return render(request, 'export.html', {'full_name': full_name})


@login_required
def detail_activity_download(request, id):
    activity = Activity.objects.get(pk=id)
    response = HttpResponse(content_type="application/ms-word")
    response["content-Disposition"] = 'attachment;filename=report' + str(datetime.datetime.now()) + '.docx'
    full_name = request.user.get_full_name()

    activityimages = ActivityImages.objects.all()
    document = Document()
    activity
    document.add_heading(activity.name, 0)
    p = document.add_paragraph(str(activity.str_user))
    p = document.add_paragraph(activity.task)
    p.add_run("_")
    p.add_run(activity.sub_task)
    p = document.add_paragraph("مدت زمان انجام فعالیت")
    p.add_run(str(activity.duration))
    p.add_run("ساعت")
    if activity.colleague:
        p = document.add_paragraph(":با مشارکت همکار / همکاران")
        p = document.add_paragraph(activity.colleague)
    p = document.add_paragraph(":تاریخ انجام فعالیت")
    p = document.add_paragraph(str(date2jalali(activity.created)))

    document.add_picture(activity.image, width=Inches(3.9))
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    for y in activityimages:
        if y.activity == activity:
            run.add_picture(y.image, width=Inches(2.3))

    p = document.add_paragraph(":توضیحات")
    p = document.add_paragraph(activity.description)
    document.add_page_break()
    document.save(response)
    return response


class PasswordsChangeView(PasswordChangeView):
    form_class = PasswordChangingForm
    success_url = reverse_lazy('password_success')
    template_name = "change_password.html"

    def get_context_data(self,  **kwargs):
        current_loggedin_user = self.request.user.get_full_name()
        context = {'full_name': current_loggedin_user}
        return context


@login_required
def password_success(request):
    full_name = request.user.get_full_name()
    return render(request, 'password_success.html', {'full_name': full_name})
